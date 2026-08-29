"""Knowledge Base Word-import: docx parsing + mapping -> module/topic apply."""
import io

from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import Client, TestCase, override_settings
from django.urls import reverse
from django_tenants.utils import get_public_schema_name, schema_context
from docx import Document

from apps.accounts.models import User
from apps.central_library.imports import (
    MappingError, apply_knowledge_base_import, build_module_payload,
)
from apps.central_library.models import (
    KnowledgeBaseImport, KnowledgeBaseModule, KnowledgeBaseTopic,
)
from shared.docx_blocks import parse_docx_blocks


def _sample_docx() -> io.BytesIO:
    d = Document()
    d.add_heading('Dashboard', level=1)
    d.add_paragraph('The dashboard is the first screen after login.')
    p = d.add_paragraph('It surfaces ')
    p.add_run('pending reviews').bold = True
    p.add_run(' and today’s appointments.')
    d.add_paragraph('Admin, Supervisor, Staff')
    d.add_heading('What users see', level=2)
    d.add_paragraph('Session review queue', style='List Bullet')
    d.add_paragraph('Notes awaiting approval', style='List Bullet')
    d.add_paragraph('Client activity trend', style='List Bullet')
    buf = io.BytesIO()
    d.save(buf)
    buf.seek(0)
    return buf


class DocxParserTests(TestCase):
    def test_extracts_headings_lists_and_inline_formatting(self):
        result = parse_docx_blocks(_sample_docx())
        blocks = result['blocks']
        kinds = [b['kind'] for b in blocks]

        self.assertEqual(kinds.count('heading'), 2)
        self.assertEqual(kinds.count('list_item'), 3)
        self.assertEqual(blocks[0]['kind'], 'heading')
        self.assertEqual(blocks[0]['level'], 1)
        self.assertIn('<strong>pending reviews</strong>', blocks[2]['html'])
        self.assertEqual(result['image_count'], 0)

    def test_rejects_non_docx(self):
        with self.assertRaises(ValueError):
            parse_docx_blocks(io.BytesIO(b'this is not a docx'))


class ApplyMappingTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_superuser(email='root@example.com', password='x')
        self.blocks = parse_docx_blocks(_sample_docx())['blocks']
        # b1 heading, b2 para, b3 para(bold), b4 para(audience), b5 heading,
        # b6/b7/b8 list items
        self.mapping = {
            'module': {
                'title': ['b1'],
                'overview': ['b2', 'b3'],
                'audience': ['b4'],
            },
            'topics': [
                {'key': 't1', 'title': ['b5'], 'items': ['b6', 'b7', 'b8']},
            ],
        }

    def _make_import(self, mapping=None):
        with schema_context(get_public_schema_name()):
            return KnowledgeBaseImport.objects.create(
                file=None, original_filename='dashboard.docx',
                blocks=self.blocks, mapping=mapping or self.mapping,
                created_by=self.user,
            )

    def test_build_payload_shapes_module_and_topics(self):
        slug, defaults, topics = build_module_payload(
            self.mapping, self.blocks, slug='', icon='bar_chart',
            display_order=10, is_active=True,
        )
        self.assertEqual(slug, 'dashboard')
        self.assertEqual(defaults['title'], 'Dashboard')
        self.assertEqual(defaults['icon'], 'bar_chart')
        self.assertEqual(defaults['audience'], ['Admin', 'Supervisor', 'Staff'])
        self.assertIn('**pending reviews**', defaults['overview'])
        self.assertEqual(len(topics), 1)
        self.assertEqual(topics[0]['title'], 'What users see')
        self.assertEqual(topics[0]['items'], [
            'Session review queue', 'Notes awaiting approval', 'Client activity trend',
        ])

    def test_apply_creates_module_with_topics(self):
        kb_import = self._make_import()
        with schema_context(get_public_schema_name()):
            module = apply_knowledge_base_import(
                kb_import, user=self.user, slug='', icon='bar_chart',
                display_order=5, is_active=True, replace_topics=True,
            )
            module.refresh_from_db()
            kb_import.refresh_from_db()

            self.assertEqual(module.slug, 'dashboard')
            self.assertEqual(module.display_order, 5)
            self.assertEqual(module.topics.count(), 1)
            self.assertEqual(module.topics.get().items[0], 'Session review queue')
            self.assertEqual(kb_import.status, KnowledgeBaseImport.Status.APPLIED)
            self.assertEqual(kb_import.target_module_id, module.id)

    def test_apply_is_idempotent_on_slug(self):
        with schema_context(get_public_schema_name()):
            apply_knowledge_base_import(
                self._make_import(), user=self.user, slug='dashboard', icon='book',
                display_order=0, is_active=True, replace_topics=True,
            )
            apply_knowledge_base_import(
                self._make_import(), user=self.user, slug='dashboard', icon='book',
                display_order=0, is_active=True, replace_topics=True,
            )
            self.assertEqual(KnowledgeBaseModule.objects.filter(slug='dashboard').count(), 1)
            self.assertEqual(
                KnowledgeBaseTopic.objects.filter(module__slug='dashboard').count(), 1,
            )

    def test_missing_module_title_is_rejected(self):
        bad = {'module': {'overview': ['b2']}, 'topics': [{'title': ['b5']}]}
        with self.assertRaises(MappingError):
            build_module_payload(bad, self.blocks, slug='', icon='book',
                                 display_order=0, is_active=True)

    def test_no_topic_title_is_rejected(self):
        bad = {'module': {'title': ['b1']}, 'topics': [{'items': ['b6']}]}
        with self.assertRaises(MappingError):
            build_module_payload(bad, self.blocks, slug='', icon='book',
                                 display_order=0, is_active=True)


DOCX_CT = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'


@override_settings(STORAGES={
    'default': {'BACKEND': 'django.core.files.storage.FileSystemStorage'},
    'staticfiles': {'BACKEND': 'django.contrib.staticfiles.storage.StaticFilesStorage'},
})
class AdminDocxImportFlowTests(TestCase):
    def setUp(self):
        self.superuser = User.objects.create_superuser(email='root@example.com', password='pw')
        self.staff = User.objects.create_user(
            email='staff@example.com', password='pw', is_staff=True, role=User.Role.ADMIN,
        )
        self.client = Client()

    def _upload(self):
        self.client.force_login(self.superuser)
        blob = _sample_docx().getvalue()
        upload = SimpleUploadedFile('dashboard.docx', blob, content_type=DOCX_CT)
        return self.client.post(
            reverse('admin:central_library_knowledgebasemodule_import_docx'),
            {'file': upload},
        )

    def test_non_superuser_is_forbidden(self):
        self.client.force_login(self.staff)
        resp = self.client.get(reverse('admin:central_library_knowledgebasemodule_import_docx'))
        self.assertEqual(resp.status_code, 403)

    def test_upload_creates_draft_and_redirects_to_map(self):
        resp = self._upload()
        kb_import = KnowledgeBaseImport.objects.get()
        self.assertEqual(resp.status_code, 302)
        self.assertIn(f'/import-docx/{kb_import.id}/', resp['Location'])
        self.assertEqual(kb_import.status, KnowledgeBaseImport.Status.DRAFT)
        self.assertTrue(kb_import.blocks)
        # the mapping screen renders
        page = self.client.get(resp['Location'])
        self.assertEqual(page.status_code, 200)
        self.assertContains(page, 'Module settings')
        self.assertContains(page, 'assign_b1')

    def test_upload_form_page_renders(self):
        self.client.force_login(self.superuser)
        resp = self.client.get(reverse('admin:central_library_knowledgebasemodule_import_docx'))
        self.assertEqual(resp.status_code, 200)
        self.assertContains(resp, 'How it works')

    def test_map_then_apply_creates_module(self):
        self._upload()
        kb_import = KnowledgeBaseImport.objects.get()
        url = reverse(
            'admin:central_library_knowledgebasemodule_map_docx',
            kwargs={'import_id': kb_import.id},
        )
        # b1 heading -> module.title, b2/b3 -> overview, b4 -> audience,
        # b5 heading -> topic 1 title, b6/b7/b8 bullets -> topic 1 items
        resp = self.client.post(url, {
            'action': 'apply',
            'slug': '',
            'icon': 'bar_chart',
            'display_order': '7',
            'is_active': 'on',
            'replace_topics': 'on',
            'topic_count': '3',
            'assign_b1': 'module.title',
            'assign_b2': 'module.overview',
            'assign_b3': 'module.overview',
            'assign_b4': 'module.audience',
            'assign_b5': 'topic.1.title',
            'assign_b6': 'topic.1.items',
            'assign_b7': 'topic.1.items',
            'assign_b8': 'topic.1.items',
        })
        self.assertEqual(resp.status_code, 302)
        module = KnowledgeBaseModule.objects.get(slug='dashboard')
        self.assertEqual(module.title, 'Dashboard')
        self.assertEqual(module.icon, 'bar_chart')
        self.assertEqual(module.display_order, 7)
        self.assertEqual(module.audience, ['Admin', 'Supervisor', 'Staff'])
        topic = module.topics.get()
        self.assertEqual(topic.title, 'What users see')
        self.assertEqual(topic.items, [
            'Session review queue', 'Notes awaiting approval', 'Client activity trend',
        ])
        kb_import.refresh_from_db()
        self.assertEqual(kb_import.status, KnowledgeBaseImport.Status.APPLIED)

    def test_preview_reports_mapping_error_without_creating(self):
        self._upload()
        kb_import = KnowledgeBaseImport.objects.get()
        url = reverse(
            'admin:central_library_knowledgebasemodule_map_docx',
            kwargs={'import_id': kb_import.id},
        )
        resp = self.client.post(url, {
            'action': 'preview', 'slug': '', 'icon': 'book',
            'display_order': '0', 'topic_count': '3',
            'assign_b2': 'module.overview',  # no title assigned
        })
        self.assertEqual(resp.status_code, 200)
        self.assertContains(resp, 'module title')
        self.assertFalse(KnowledgeBaseModule.objects.filter(slug='dashboard').exists())
